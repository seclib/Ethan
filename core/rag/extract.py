"""Core text-extraction facade for file upload -> RAG ingestion.

Design rules (AGENTS.md / architecture plan section 1.4 #6):
- Extraction logic lives in Core -- the WebUI only calls ``extract_text``
  through the API gateway.
- Pure-Python fallbacks keep zero mandatory heavy deps.  Optional engines
  (PyPDF, python-docx) are imported lazily and degrade gracefully.
- Extensible registry: new formats can register a ``TextExtractor`` without
  touching existing call-sites (open/closed).
"""

from __future__ import annotations

import csv
import io
import json
import logging
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Callable

logger = logging.getLogger(__name__)

# -- Types ------------------------------------------------------------------


@dataclass
class ExtractedText:
    """Result of extracting text from a file.

    Attributes:
        text: Full extracted text content.
        page_count: Number of pages (if detectable, else 0).
        metadata: Engine/format metadata for diagnostics.
    """

    text: str
    page_count: int = 0
    metadata: dict[str, Any] = field(default_factory=dict)


# -- Optional engine detection -----------------------------------------------

try:  # noqa: SIM110
    import pypdf  # type: ignore[import-untyped]
    _HAS_PYPDF = True
    logger.debug("pypdf available for PDF extraction")
except ImportError:  # pragma: no cover
    _HAS_PYPDF = False
    logger.debug("pypdf not available; PDF extraction will raise informative error")

try:
    import docx  # type: ignore[import-untyped]
    _HAS_DOCX = True
    logger.debug("python-docx available for DOCX extraction")
except ImportError:  # pragma: no cover
    _HAS_DOCX = False
    logger.debug("python-docx not available; DOCX extraction disabled")


# -- HTML text extraction (stdlib) -------------------------------------------


class _HTMLTextExtractor(HTMLParser):
    """Strip HTML tags, preserve text."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._parts: list[str] = []

    def handle_data(self, data: str) -> None:
        self._parts.append(data)

    @property
    def text(self) -> str:
        return "".join(self._parts)


# -- Extractor registry ------------------------------------------------------

ExtractorFn = Callable[[bytes, str], ExtractedText]

_extractors: dict[str, ExtractorFn] = {}


def register_extractor(extension: str, fn: ExtractorFn) -> None:
    """Register a text extractor for a given file extension (lowercase, no dot)."""
    _extractors[extension.lower()] = fn


def get_extractor(extension: str) -> ExtractorFn | None:
    """Return the registered extractor for an extension, or None."""
    return _extractors.get(extension.lower().lstrip("."))


# -- Built-in extractors -----------------------------------------------------


def _extract_plain(data: bytes, filename: str) -> ExtractedText:
    """Plain-text extraction (UTF-8 with fallback)."""
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="replace")
    return ExtractedText(text=text, metadata={"engine": "builtin-utf8"})


def _extract_csv(data: bytes, filename: str) -> ExtractedText:
    """CSV -> text (one row per line, tab-separated)."""
    try:
        text = data.decode("utf-8-sig")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = ["\t".join(row) for row in reader if row]
    return ExtractedText(
        text="\n".join(rows),
        page_count=len(rows),
        metadata={"engine": "builtin-csv", "rows": len(rows)},
    )


def _extract_json(data: bytes, filename: str) -> ExtractedText:
    """JSON -> pretty-printed text for chunking."""
    try:
        text = data.decode("utf-8")
        obj = json.loads(text)
        rendered = json.dumps(obj, indent=2, ensure_ascii=False)
    except (UnicodeDecodeError, json.JSONDecodeError):
        rendered = data.decode("latin-1", errors="replace")
    return ExtractedText(text=rendered, metadata={"engine": "builtin-json"})


def _extract_html(data: bytes, filename: str) -> ExtractedText:
    """HTML -> plain text (strips tags)."""
    try:
        raw = data.decode("utf-8")
    except UnicodeDecodeError:
        raw = data.decode("latin-1", errors="replace")
    parser = _HTMLTextExtractor()
    parser.feed(raw)
    return ExtractedText(text=parser.text.strip(), metadata={"engine": "builtin-html"})


def _extract_pdf(data: bytes, filename: str) -> ExtractedText:
    """PDF extraction via pypdf (if available)."""
    if not _HAS_PYPDF:
        raise ValueError(
            "PDF extraction requires 'pypdf' -- install with: pip install pypdf "
            "(or use the 'marker'/'docling' extraction engine)"
        )
    reader = pypdf.PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    text = "\n".join(pages)
    return ExtractedText(
        text=text,
        page_count=len(reader.pages),
        metadata={"engine": "pypdf", "pages": len(reader.pages)},
    )


def _extract_docx(data: bytes, filename: str) -> ExtractedText:
    """DOCX extraction via python-docx (if available)."""
    if not _HAS_DOCX:
        raise ValueError(
            "DOCX extraction requires 'python-docx' -- install with: pip install python-docx"
        )
    doc = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_text: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                tables_text.append(" | ".join(cells))
    text = "\n".join(paragraphs + tables_text)
    return ExtractedText(
        text=text,
        page_count=len(doc.paragraphs),
        metadata={"engine": "python-docx", "paragraphs": len(paragraphs)},
    )


# -- Auto-register built-in extractors ---------------------------------------

register_extractor("txt", _extract_plain)
register_extractor("md", _extract_plain)
register_extractor("markdown", _extract_plain)
register_extractor("ts", _extract_plain)
register_extractor("tsx", _extract_plain)
register_extractor("js", _extract_plain)
register_extractor("jsx", _extract_plain)
register_extractor("py", _extract_plain)
register_extractor("json", _extract_json)
register_extractor("csv", _extract_csv)
register_extractor("html", _extract_html)
register_extractor("htm", _extract_html)
register_extractor("pdf", _extract_pdf)
register_extractor("docx", _extract_docx)


# -- Public API --------------------------------------------------------------


def extract_text(
    content: bytes,
    filename: str | Path = "",
    content_type: str = "",
) -> ExtractedText:
    """Extract text from raw file content.

    Args:
        content: Raw file bytes.
        filename: Filename (used to determine extension).
        content_type: MIME type (fallback if extension unknown).

    Returns:
        ExtractedText with full text + metadata.

    Raises:
        ValueError: If extraction fails for a known-but-unsupported format.
    """
    name = str(filename)
    ext = Path(name).suffix.lower().lstrip(".")

    # Fallback: sniff content type if no extension
    if not ext and content_type:
        ext = content_type.split("/")[-1].split("+")[-1]

    extractor = get_extractor(ext)
    if extractor is None:
        # Fall back to plain text for unknown types
        logger.warning("No extractor for .%s -- falling back to plain text", ext)
        return _extract_plain(content, name)

    return extractor(content, name)</arg_value></tool_call>